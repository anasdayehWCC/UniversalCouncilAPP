import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt

from common.database.postgres_models import MinuteVersion


def _strip_html(html: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    return text


def _extract_citations(text: str) -> list[str]:
    matches = re.findall(r"\[(\d+(?:-\d+)?)\]", text)
    return list(dict.fromkeys(matches))


def _add_cover(document: Document, minute_version: MinuteVersion) -> None:
    minute = minute_version.minute
    document.add_heading("Social Care Minute", level=0)
    meta_table = document.add_table(rows=0, cols=2)
    meta_table.style = "LightShading-Accent1"

    def add_row(label: str, value: str | None):
        row_cells = meta_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value or "-"

    add_row("Case reference", minute.case_reference)
    add_row("Visit type", minute.visit_type)
    add_row("Service domain", getattr(minute, "service_domain_id", None) and str(minute.service_domain_id))
    add_row("Created", minute.created_datetime.isoformat() if minute.created_datetime else "-")
    add_row("Updated", minute.updated_datetime.isoformat() if minute.updated_datetime else "-")


def build_docx(minute_version: MinuteVersion, output_path: Path) -> Path:
    document = Document()
    minute = minute_version.minute
    _add_cover(document, minute_version)
    document.add_heading(minute.template_name, level=1)

    text_body = _strip_html(minute_version.html_content or "")
    paragraphs = [p.strip() for p in text_body.split("\n") if p.strip()]
    for para in paragraphs:
        p = document.add_paragraph(para)
        p.style.font.size = Pt(11)

    citations = _extract_citations(text_body)
    if citations:
        document.add_heading("Citations", level=2)
        for c in citations:
            document.add_paragraph(f"[{c}] – see transcript timeline")

    document.save(output_path)
    return output_path
